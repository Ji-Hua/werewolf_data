import json
import re

import docx
import pandas as pd

import codecs


from engine_factory import EngineFactory


CHINESE_NUMBER_DICT = {
    1: '一',
    2: '二',
    3: '三',
    4: '四',
    5: '五',
    6: '六',
    7: '七',
    8: '八',
    9: '九'
}


def read_table(doc_name, table_id: int):
    # Load the first table from your document. In your example file,
    # there is only one table, so I just grab the first one.
    document = docx.Document(doc_name)
    table = document.tables[table_id]

    # Data will be a list of rows represented as dictionaries
    # containing each row's data.
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    
    return data

def get_game_name_by_id(doc_name, game_id: int):
    # get the page of game id
    document = docx.Document(doc_name)
    paragraph_id = None
    for i, p in enumerate(document.paragraphs):
        pattern = re.compile(f'第{CHINESE_NUMBER_DICT[game_id]}局')
        if re.search(pattern, p.text):
            paragraph_id = i + 1
            break
    else:
        raise ValueError(f"{game_id} out of range")
    template = document.paragraphs[paragraph_id].text.split('：')[-1]
    return re.sub('vs\.?', 'vs.', template)

def get_game_winner_by_id(doc_name, game_id: int):
    # get winner infomation
    document = docx.Document(doc_name)
    start_id, end_id = None, len(document.paragraphs)
    for i, p in enumerate(document.paragraphs):
        p_start = re.compile(f'第{CHINESE_NUMBER_DICT[game_id]}局')
        p_end = re.compile(f'第{CHINESE_NUMBER_DICT[game_id+1]}局')
        if re.search(p_start, p.text):
            start_id = i + 1
        elif re.search(p_end, p.text):
            end_id = i
            break
    winner = None
    for p in document.paragraphs[start_id:end_id]:
        pattern = re.compile(f'胜方：\s*(\S+)阵营')
        match = re.search(pattern, p.text)
        if match:
            winner = match.group(1)
            return winner
    else:
        raise ValueError(f"Winner not found of {game_id}")


def get_game_data(doc_name, game_id: int):
    vote_table_id = (game_id - 1) * 2
    action_table_id = (game_id - 1) * 2 + 1

    vote_df = pd.DataFrame(read_table(doc_name, vote_table_id))
    action_df = pd.DataFrame(read_table(doc_name, action_table_id))

    factory = EngineFactory()
    name = get_game_name_by_id(doc_name, game_id)
    winner = get_game_winner_by_id(doc_name, game_id)
    cleaned_data = {"template": name, "winner": winner}
    parser = factory.construct(name)
    if parser:
        parser.read_data(action_df, vote_df)
        cleaned_data["data"] = parser.parse(winner)
        return cleaned_data


def write_cleaned_data(cleaned_data, dest):
    with codecs.open(dest, 'w+',  encoding = "utf-8") as f:
        json.dump(cleaned_data, f, ensure_ascii=False, indent=2)



if __name__ == '__main__':
    doc_name_base = "./data/HCSSA 桌游社狼人杀游戏裁判表"
    games = [
        ("0829", 1), ("0829", 2), ("0829", 3), ("0829", 4),
        ("0905", 1), ("0905", 2), ("0905", 3),
        ("0912", 1), ("0912", 2), ("0912", 3), ("0912", 4),
        ("0919", 1), ("0919", 2), ("0919", 3), ("0919", 4),
        ("0926", 1), ("0926", 2),("0926", 3), ("0926", 4),
        ("1003", 1), ("1003", 2), ("1003", 3),
        ("1010", 1), ("1010", 2), ("1010", 3), ("1010", 4),
        ("1017", 1), ("1017", 2), ("1017", 3),
        ("1024", 1), ("1024", 2), ("1024", 3), ("1024", 4), 
        ("1031", 1), ("1031", 2),
        ("1107", 1), ("1107", 2), ("1107", 3),
        ("1114", 1), ("1114", 2), ("1114", 3), ("1114", 4),
        ("1121", 1), ("1121", 2),
        ("1127", 1), ("1127", 2), ("1127", 3), ("1127", 4)
    ]
    
    for game in games:
        print(game)
        doc_name = f"{doc_name_base} {game[0]}.docx"
        game_id = game[1]

        new_name = doc_name.split()[-1].split('.')[0]
        dest = f'./cleaned_data/{new_name}-{game_id}.json'
        cleaned_data = get_game_data(doc_name, game_id)
        if cleaned_data:
            write_cleaned_data(cleaned_data, dest)
    
    '''
    game = ("1127", 1)
    doc_name = f"{doc_name_base} {game[0]}.docx"
    game_id = game[1]
    name, winner, vote_df, action_df, cleaned_data = get_game_data(doc_name, game_id)
    new_name = doc_name.split()[-1].split('.')[0]
    dest = f'./cleaned_data/{new_name}-{game_id}.json'
    write_cleaned_data(cleaned_data, dest)
    '''
